from docx import Document

def split_docx_by_paragraphs(input_file, num_parts):
    doc = Document(input_file)
    paragraphs = doc.paragraphs
    total_paragraphs = len(paragraphs)
    chunk_size = total_paragraphs // num_parts

    for i in range(num_parts):
        new_doc = Document()
        start = i * chunk_size
        end = (i + 1) * chunk_size if i < num_parts - 1 else total_paragraphs
        for para in paragraphs[start:end]:
            new_doc.add_paragraph(para.text)
        
        new_doc.save(f'output_part_{i+1}.docx')

split_docx_by_paragraphs("blackbasta.docx", 10)
