from docx import Document

def merge_docx(output_file, input_files):
    merged_doc = Document()

    for i, file in enumerate(input_files):
        print(f"Processing file {i+1}/{len(input_files)}: {file}")
        doc = Document(file)
        
        # Append each paragraph from the current document
        for para in doc.paragraphs:
            merged_doc.add_paragraph(para.text)
        
        # Add a page break after each document (except the last one)
        if i < len(input_files) - 1:
            merged_doc.add_page_break()

    merged_doc.save(output_file)
    print(f"\nMerged document saved as: {output_file}")

# List your custom file names here (ensure they exist in the directory)
input_files = [
"output_part_1_translated.docx", "output_part_2_transalated.docx", "output_part_3_translated.docx", "output_part_4_translated.docx", "output_part_5_transalated.docx", "output_part_6_translated.docx", "output_part_7_translated.docx", "output_part_8_translated.docx", "output_part_9_translated.docx", "output_part_10_translated.docx"
]

# Merge them into a single document
merge_docx("merged_output.docx", input_files)
